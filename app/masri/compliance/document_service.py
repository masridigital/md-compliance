"""Template upload + AI document generation pipeline.

Three public entry points:

    * :func:`extract_placeholders` — scan a .docx for ``{{NAME}}`` tokens.
    * :func:`generate_from_scratch` — produce a new doc using prompts only.
    * :func:`generate_from_template` — fill template placeholders then ask
      the LLM to complete/polish the remaining prose.

All LLM traffic goes through :class:`app.masri.llm_service.LLMService` so
per-tier routing, adapter normalization, and budget/rate-limit
enforcement come for free.

Output bytes land in the storage router under ``role="reports"``; the
returned ``ComplianceDocumentVersion`` carries the resulting
``storage_key``.
"""

from __future__ import annotations

import io
import logging
import re
from datetime import datetime
from typing import Any

from app import db
from app.masri.compliance import framework_meta
from app.masri.compliance.service import get_exemption_profile
from app.masri.new_models import (
    ComplianceDocument,
    ComplianceDocumentVersion,
    DocumentTemplate,
    Questionnaire,
)

logger = logging.getLogger(__name__)


PLACEHOLDER_RE = re.compile(r"\{\{\s*([A-Za-z0-9_\.]+)\s*\}\}")


# ── Placeholder extraction ────────────────────────────────────────────────

def extract_placeholders(docx_bytes: bytes) -> list[str]:
    """Return the unique placeholder names found in a .docx file."""
    from docx import Document as DocxDocument

    doc = DocxDocument(io.BytesIO(docx_bytes))
    found: set[str] = set()

    for para in doc.paragraphs:
        for match in PLACEHOLDER_RE.findall(para.text or ""):
            found.add(match)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for match in PLACEHOLDER_RE.findall(para.text or ""):
                        found.add(match)
    return sorted(found)


def build_context(tenant_id: str, framework_slug: str | None = None) -> dict[str, Any]:
    """Assemble the canonical placeholder context for a tenant.

    Answers come from the active questionnaire (if any). Org data is
    pulled from the :class:`Tenant`. Anything the placeholder map wires
    to a ``questionnaire.*`` or ``org.*`` path resolves through this
    context.
    """
    from app.models import Tenant

    tenant = db.session.get(Tenant, tenant_id)
    org: dict[str, Any] = {}
    if tenant:
        org = {
            "name": tenant.name,
            "contact_email": getattr(tenant, "contact_email", None),
            "contact_name": getattr(tenant, "contact_name", None),
        }

    answers: dict[str, Any] = {}
    if framework_slug:
        q = (
            db.session.execute(
                db.select(Questionnaire)
                .filter(Questionnaire.tenant_id == tenant_id)
                .filter(Questionnaire.framework_slug == framework_slug)
                .filter(Questionnaire.status != "archived")
                .order_by(Questionnaire.date_added.desc())
            )
            .scalars()
            .first()
        )
        if q:
            answers = q.answers or {}

    return {
        "org": org,
        "questionnaire": answers,
        "generated_at": datetime.utcnow().isoformat(),
    }


def _resolve_path(path: str, context: dict[str, Any]) -> str | None:
    """Walk a dotted path through the context. Return None when unknown."""
    cur: Any = context
    for part in path.split("."):
        if isinstance(cur, dict) and part in cur:
            cur = cur[part]
        else:
            return None
    if cur is None:
        return None
    return str(cur)


def _replace_in_paragraph(paragraph, mapping: dict[str, str]) -> None:
    """Replace placeholders in a single paragraph, preserving run style.

    python-docx splits text across runs; naive ``para.text = ...``
    destroys formatting. Strategy: join runs, run regex, then rewrite
    first run with the replacement and blank the rest. This loses
    inline formatting _within_ a replaced placeholder but keeps
    paragraph-level styling. Placeholders are plain-text tokens so this
    is acceptable.
    """
    if not paragraph.runs:
        return
    full = "".join(run.text or "" for run in paragraph.runs)
    replaced = full
    for key, value in mapping.items():
        replaced = replaced.replace("{{" + key + "}}", value)
        replaced = replaced.replace("{{ " + key + " }}", value)
    if replaced != full:
        paragraph.runs[0].text = replaced
        for run in paragraph.runs[1:]:
            run.text = ""


def inject_known_values(
    docx_bytes: bytes,
    placeholder_map: dict[str, str],
    context: dict[str, Any],
) -> tuple[bytes, str, list[str]]:
    """Replace mapped placeholders and return (new_bytes, plain_text, unmapped).

    ``placeholder_map`` — ``{placeholder_name: path}`` where path is either
    a dotted context path (``org.name``, ``questionnaire.employee_count``)
    or the sentinel ``"leave_for_ai"`` / ``"manual_input"``.
    """
    from docx import Document as DocxDocument

    doc = DocxDocument(io.BytesIO(docx_bytes))

    literals: dict[str, str] = {}
    unmapped: list[str] = []
    for placeholder, path in (placeholder_map or {}).items():
        if not path or path in ("leave_for_ai", "manual_input"):
            unmapped.append(placeholder)
            continue
        resolved = _resolve_path(path, context)
        if resolved is None:
            unmapped.append(placeholder)
            continue
        literals[placeholder] = resolved

    for para in doc.paragraphs:
        _replace_in_paragraph(para, literals)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_paragraph(para, literals)

    out = io.BytesIO()
    doc.save(out)
    new_bytes = out.getvalue()

    # Plain-text dump for the LLM so it sees the partially filled doc.
    text_doc = DocxDocument(io.BytesIO(new_bytes))
    lines = [p.text for p in text_doc.paragraphs]
    for table in text_doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text:
                        lines.append(para.text)
    plain_text = "\n".join(lines)

    return new_bytes, plain_text, unmapped


# ── LLM generation ────────────────────────────────────────────────────────

def _llm_complete(system_prompt: str, user_prompt: str) -> str:
    """Route to the configured Tier 4 model via LLMService."""
    from app.masri.llm_service import LLMService

    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_prompt},
    ]
    result = LLMService.chat(messages, feature="policy_draft", max_tokens=6000)
    return (result or {}).get("content", "")


def _markdown_to_docx(markdown_text: str, title: str) -> bytes:
    """Render AI markdown output into a basic styled .docx.

    Kept minimal — one Heading 1 (title), Heading 2 for ``##`` lines,
    Normal for everything else, and a signature block appended. Good
    enough for a draft; designers can upload branded templates to get
    fancier output.
    """
    from docx import Document as DocxDocument

    doc = DocxDocument()
    doc.add_heading(title, level=1)

    for raw in (markdown_text or "").splitlines():
        line = raw.rstrip()
        if not line.strip():
            doc.add_paragraph("")
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.lstrip().startswith(("- ", "* ")):
            doc.add_paragraph(line.lstrip()[2:], style="List Bullet")
        else:
            doc.add_paragraph(line)

    doc.add_paragraph("")
    doc.add_paragraph("__________________________")
    doc.add_paragraph("Approved by")
    doc.add_paragraph(f"Date: {datetime.utcnow().strftime('%Y-%m-%d')}")

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def generate_from_scratch(
    *,
    tenant_id: str,
    doc_type: str,
    framework_slug: str | None,
    title: str | None = None,
    project_id: str | None = None,
    user_id: str | None = None,
) -> ComplianceDocument:
    """Generate a brand-new .docx using the prompt registry."""
    from app.masri.compliance import prompts

    context = build_context(tenant_id, framework_slug)
    system_prompt = prompts.get_system_prompt(doc_type, framework_slug)
    user_prompt = prompts.build_user_prompt(
        doc_type=doc_type,
        framework_slug=framework_slug,
        org_profile=context["org"],
        answers=context["questionnaire"],
        exemption_profile=_exemption_snapshot(tenant_id, framework_slug),
    )

    ai_text = _llm_complete(system_prompt, user_prompt)
    final_title = title or prompts.default_title(doc_type, context["org"])
    docx_bytes = _markdown_to_docx(ai_text, final_title)

    return _persist(
        tenant_id=tenant_id,
        project_id=project_id,
        framework_slug=framework_slug,
        doc_type=doc_type,
        title=final_title,
        template_id=None,
        docx_bytes=docx_bytes,
        content_text=ai_text,
        prompt_used=system_prompt,
        mode="from_scratch",
        user_id=user_id,
    )


def generate_from_template(
    *,
    tenant_id: str,
    template_id: str,
    title: str | None = None,
    project_id: str | None = None,
    user_id: str | None = None,
) -> ComplianceDocument:
    from app.masri.compliance import prompts
    from app.masri.storage_router import get_file

    template = db.session.get(DocumentTemplate, template_id)
    if not template:
        raise ValueError(f"Template not found: {template_id}")
    if template.tenant_id and template.tenant_id != tenant_id and not template.is_global:
        raise PermissionError("Template belongs to another tenant")

    docx_bytes = get_file(template.storage_key, role="reports")
    if not docx_bytes:
        raise ValueError(f"Template file missing from storage: {template.storage_key}")

    context = build_context(tenant_id, template.framework_slug)
    filled_bytes, plain_text, unmapped = inject_known_values(
        docx_bytes, template.placeholder_map or {}, context
    )

    system_prompt = prompts.get_edit_prompt(template.doc_type, template.framework_slug)
    user_prompt = prompts.build_edit_user_prompt(
        doc_type=template.doc_type,
        framework_slug=template.framework_slug,
        org_profile=context["org"],
        answers=context["questionnaire"],
        exemption_profile=_exemption_snapshot(tenant_id, template.framework_slug),
        template_text=plain_text,
        unmapped=unmapped,
    )
    ai_text = _llm_complete(system_prompt, user_prompt)

    # Re-render the AI-completed version. We rebuild the docx rather than
    # trying to splice LLM output back into the template because Word
    # paragraph IDs don't survive round-trips.
    final_title = title or template.name
    docx_out = _markdown_to_docx(ai_text, final_title)

    return _persist(
        tenant_id=tenant_id,
        project_id=project_id,
        framework_slug=template.framework_slug,
        doc_type=template.doc_type,
        title=final_title,
        template_id=template.id,
        docx_bytes=docx_out,
        content_text=ai_text,
        prompt_used=system_prompt,
        mode="from_template",
        user_id=user_id,
    )


def upload_template(
    *,
    tenant_id: str | None,
    docx_bytes: bytes,
    file_name: str,
    doc_type: str,
    name: str | None = None,
    description: str | None = None,
    framework_slug: str | None = None,
    is_global: bool = False,
    user_id: str | None = None,
) -> DocumentTemplate:
    """Persist an uploaded .docx template and auto-extract its placeholders."""
    from app.masri.storage_router import store_file

    placeholders = extract_placeholders(docx_bytes)
    folder = (
        f"compliance/templates/global"
        if is_global or tenant_id is None
        else f"compliance/templates/{tenant_id}"
    )
    storage_key = store_file(
        file_data=docx_bytes,
        file_name=file_name,
        folder=folder,
        role="reports",
        tenant_id=tenant_id,
    )
    template = DocumentTemplate(
        tenant_id=None if is_global else tenant_id,
        name=name or file_name,
        description=description,
        framework_slug=framework_slug,
        doc_type=doc_type,
        storage_key=storage_key,
        placeholders=placeholders,
        placeholder_map={p: "leave_for_ai" for p in placeholders},
        is_global=is_global,
        created_by_user_id=user_id,
    )
    db.session.add(template)
    db.session.commit()
    return template


# ── Persistence helpers ────────────────────────────────────────────────────

def _exemption_snapshot(
    tenant_id: str, framework_slug: str | None
) -> dict[str, Any] | None:
    if not framework_slug:
        return None
    profile = get_exemption_profile(tenant_id, framework_slug)
    return profile.as_dict() if profile else None


def _persist(
    *,
    tenant_id: str,
    project_id: str | None,
    framework_slug: str | None,
    doc_type: str,
    title: str,
    template_id: str | None,
    docx_bytes: bytes,
    content_text: str,
    prompt_used: str,
    mode: str,
    user_id: str | None,
) -> ComplianceDocument:
    """Upload bytes, record document + version."""
    from app.masri.storage_router import store_file

    folder = f"compliance/documents/{tenant_id}"
    safe_stub = re.sub(r"[^a-zA-Z0-9_.-]+", "-", title.lower())[:80]
    file_name = f"{safe_stub}-v1.docx"
    storage_key = store_file(
        file_data=docx_bytes,
        file_name=file_name,
        folder=folder,
        role="reports",
        tenant_id=tenant_id,
    )

    document = ComplianceDocument(
        tenant_id=tenant_id,
        project_id=project_id,
        framework_slug=framework_slug,
        doc_type=doc_type,
        title=title,
        template_id=template_id,
        current_version=1,
        created_by_user_id=user_id,
    )
    db.session.add(document)
    db.session.flush()

    version = ComplianceDocumentVersion(
        document_id=document.id,
        version_num=1,
        storage_key=storage_key,
        content_text=content_text,
        prompt_used=prompt_used,
        generation_mode=mode,
        generated_by_user_id=user_id,
        meta={"framework_slug": framework_slug},
    )
    db.session.add(version)
    db.session.commit()
    return document


def list_documents(
    tenant_id: str,
    *,
    framework_slug: str | None = None,
) -> list[ComplianceDocument]:
    stmt = db.select(ComplianceDocument).filter(
        ComplianceDocument.tenant_id == tenant_id
    )
    if framework_slug:
        stmt = stmt.filter(ComplianceDocument.framework_slug == framework_slug)
    stmt = stmt.order_by(ComplianceDocument.date_added.desc())
    return db.session.execute(stmt).scalars().all()


def list_templates(tenant_id: str | None) -> list[DocumentTemplate]:
    stmt = db.select(DocumentTemplate).filter(
        db.or_(
            DocumentTemplate.is_global.is_(True),
            DocumentTemplate.tenant_id == tenant_id,
        )
    ).order_by(DocumentTemplate.date_added.desc())
    return db.session.execute(stmt).scalars().all()


def doc_types_for(framework_slug: str) -> list[str]:
    meta = framework_meta.load(framework_slug) or {}
    types: set[str] = set()
    for section in (meta.get("sections") or {}).values():
        for dt in section.get("doc_types") or []:
            types.add(dt)
    return sorted(types)
