"""
WISP Export — generates branded PDF and DOCX from WISPDocument model.
Uses reportlab for PDF, python-docx for DOCX.
Falls back to HTML if libraries unavailable.
"""

import json
import logging
import os
from datetime import datetime

logger = logging.getLogger(__name__)

SECTIONS = [
    ("firm_profile", "1. Organization Overview"),
    ("qualified_individual", "2. Qualified Individual"),
    ("asset_inventory", "3. Information Asset Inventory"),
    ("risk_assessment", "4. Risk Assessment"),
    ("access_controls", "5. Access Control Policy"),
    ("encryption", "6. Encryption & Data Protection"),
    ("monitoring", "7. Monitoring & Audit Logging"),
    ("incident_response", "8. Incident Response Plan"),
    ("training", "9. Employee Training Program"),
    ("vendor_management", "10. Third-Party Service Provider Management"),
    ("physical_security", "11. Physical Security Policy"),
    ("annual_review", "12. Annual Review & Board Reporting"),
]

# Map section keys to WISPDocument JSON columns
_SECTION_FIELD_MAP = {
    "firm_profile": None,  # built from scalar fields
    "qualified_individual": None,  # built from scalar fields
    "asset_inventory": "asset_inventory_json",
    "risk_assessment": "risk_assessment_json",
    "access_controls": "access_control_answers_json",
    "encryption": "encryption_answers_json",
    "monitoring": None,
    "incident_response": "incident_response_json",
    "training": "training_program_json",
    "vendor_management": "third_party_vendors_json",
    "physical_security": "physical_security_json",
    "annual_review": "annual_review_json",
}


class WISPExporter:
    def __init__(self, wisp, branding: dict):
        """
        Args:
            wisp: WISPDocument model instance
            branding: dict with app_name, logo_url, primary_color, support_email
        """
        self.wisp = wisp
        self.branding = branding or {}

    def _get_section_content(self, section_key: str) -> str:
        """Pull from wisp.generated_text_json[section_key] or build from raw answers."""
        # First try LLM-generated text
        generated = self.wisp.generated_text_json
        if generated:
            if isinstance(generated, str):
                try:
                    generated = json.loads(generated)
                except (json.JSONDecodeError, TypeError):
                    generated = {}
            if isinstance(generated, dict) and section_key in generated:
                return generated[section_key]

        # Fall back to raw wizard data
        if section_key == "firm_profile":
            parts = []
            if self.wisp.firm_name:
                parts.append(f"Organization: {self.wisp.firm_name}")
            if self.wisp.firm_type:
                parts.append(f"Firm Type: {self.wisp.firm_type.replace('_', ' ').title()}")
            if self.wisp.state_of_incorporation:
                parts.append(f"State: {self.wisp.state_of_incorporation}")
            if self.wisp.employee_count_range:
                parts.append(f"Employees: {self.wisp.employee_count_range}")
            if self.wisp.client_record_count_range:
                parts.append(f"Client Records: {self.wisp.client_record_count_range}")
            return "\n".join(parts) if parts else "(Not provided)"

        if section_key == "qualified_individual":
            parts = []
            if self.wisp.qi_name:
                parts.append(f"Name: {self.wisp.qi_name}")
            if self.wisp.qi_title:
                parts.append(f"Title: {self.wisp.qi_title}")
            if self.wisp.qi_email:
                parts.append(f"Email: {self.wisp.qi_email}")
            if self.wisp.qi_is_third_party:
                parts.append("Third-party QI: Yes")
            return "\n".join(parts) if parts else "(Not provided)"

        field = _SECTION_FIELD_MAP.get(section_key)
        if field:
            raw = getattr(self.wisp, field, None)
            if raw:
                if isinstance(raw, str):
                    try:
                        raw = json.loads(raw)
                    except (json.JSONDecodeError, TypeError):
                        return raw
                if isinstance(raw, dict):
                    return "\n".join(
                        f"{k.replace('_', ' ').title()}: {v}" for k, v in raw.items()
                    )
                if isinstance(raw, list):
                    return "\n".join(f"- {item}" for item in raw)
                return str(raw)

        return "(Content not yet generated for this section)"

    def export_pdf(self, output_path: str) -> str:
        """
        Generate branded PDF using reportlab.
        Falls back gracefully if reportlab not installed.
        """
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.colors import HexColor
            from reportlab.lib.units import inch
            from reportlab.platypus import (
                SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle,
            )
        except ImportError:
            logger.warning("reportlab not installed, falling back to HTML export")
            return self._export_html(output_path.replace(".pdf", ".html"))

        brand_color = HexColor(self.branding.get("primary_color", "#0066CC"))
        app_name = self.branding.get("app_name", "Masri Digital")

        doc = SimpleDocTemplate(
            output_path,
            pagesize=letter,
            topMargin=0.75 * inch,
            bottomMargin=0.75 * inch,
            leftMargin=1 * inch,
            rightMargin=1 * inch,
        )

        styles = getSampleStyleSheet()
        styles.add(ParagraphStyle(
            "CoverTitle",
            parent=styles["Title"],
            fontSize=26,
            spaceAfter=12,
            textColor=brand_color,
        ))
        styles.add(ParagraphStyle(
            "CoverSubtitle",
            parent=styles["Normal"],
            fontSize=14,
            spaceAfter=6,
            textColor=HexColor("#6E6E73"),
        ))
        styles.add(ParagraphStyle(
            "SectionHeading",
            parent=styles["Heading1"],
            fontSize=16,
            spaceBefore=24,
            spaceAfter=12,
            textColor=brand_color,
        ))
        styles.add(ParagraphStyle(
            "BodyText",
            parent=styles["Normal"],
            fontSize=11,
            leading=15,
            spaceAfter=8,
        ))
        styles.add(ParagraphStyle(
            "Footer",
            parent=styles["Normal"],
            fontSize=9,
            textColor=HexColor("#8E8E93"),
        ))

        elements = []

        # Cover page
        elements.append(Spacer(1, 2 * inch))
        if self.wisp.firm_name:
            elements.append(Paragraph(self.wisp.firm_name, styles["CoverTitle"]))
        elements.append(Paragraph("Written Information Security Program", styles["CoverTitle"]))
        elements.append(Spacer(1, 0.5 * inch))
        elements.append(Paragraph(f"Prepared by {app_name}", styles["CoverSubtitle"]))
        elements.append(Paragraph(
            f"Date: {datetime.utcnow().strftime('%B %d, %Y')}", styles["CoverSubtitle"]
        ))
        if self.wisp.qi_name:
            elements.append(Paragraph(
                f"Qualified Individual: {self.wisp.qi_name}"
                + (f", {self.wisp.qi_title}" if self.wisp.qi_title else ""),
                styles["CoverSubtitle"],
            ))
        elements.append(Spacer(1, 1 * inch))
        elements.append(Paragraph("CONFIDENTIAL", styles["CoverSubtitle"]))
        elements.append(PageBreak())

        # Sections
        for section_key, section_title in SECTIONS:
            content = self._get_section_content(section_key)
            elements.append(Paragraph(section_title, styles["SectionHeading"]))
            for line in content.split("\n"):
                line = line.strip()
                if line:
                    # Escape XML special chars for reportlab
                    line = (line.replace("&", "&amp;").replace("<", "&lt;")
                            .replace(">", "&gt;"))
                    elements.append(Paragraph(line, styles["BodyText"]))
            elements.append(Spacer(1, 0.25 * inch))

        # Footer on each page
        def _add_footer(canvas, doc):
            canvas.saveState()
            footer = f"Prepared by {app_name} | masridigital.com | Confidential"
            canvas.setFont("Helvetica", 8)
            canvas.setFillColor(HexColor("#8E8E93"))
            canvas.drawString(inch, 0.5 * inch, footer)
            canvas.drawRightString(
                letter[0] - inch, 0.5 * inch, f"Page {doc.page}"
            )
            canvas.restoreState()

        doc.build(elements, onFirstPage=_add_footer, onLaterPages=_add_footer)
        logger.info("WISP PDF exported to %s", output_path)
        return output_path

    def export_docx(self, output_path: str) -> str:
        """
        Generate editable DOCX using python-docx.
        Falls back gracefully if python-docx not installed.
        """
        try:
            from docx import Document
            from docx.shared import Pt, Inches, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            logger.warning("python-docx not installed, falling back to HTML export")
            return self._export_html(output_path.replace(".docx", ".html"))

        brand_hex = self.branding.get("primary_color", "#0066CC").lstrip("#")
        brand_rgb = RGBColor(
            int(brand_hex[0:2], 16), int(brand_hex[2:4], 16), int(brand_hex[4:6], 16)
        )
        app_name = self.branding.get("app_name", "Masri Digital")

        doc = Document()

        # Set default font
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)

        # Cover page
        cover = doc.add_paragraph()
        cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if self.wisp.firm_name:
            run = cover.add_run(self.wisp.firm_name + "\n")
            run.bold = True
            run.font.size = Pt(24)
            run.font.color.rgb = brand_rgb

        run = cover.add_run("Written Information Security Program\n\n")
        run.bold = True
        run.font.size = Pt(18)

        run = cover.add_run(f"Prepared by {app_name}\n")
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x6E, 0x6E, 0x73)

        run = cover.add_run(f"Date: {datetime.utcnow().strftime('%B %d, %Y')}\n")
        run.font.size = Pt(12)

        if self.wisp.qi_name:
            qi_text = f"Qualified Individual: {self.wisp.qi_name}"
            if self.wisp.qi_title:
                qi_text += f", {self.wisp.qi_title}"
            run = cover.add_run(qi_text + "\n")
            run.font.size = Pt(12)

        run = cover.add_run("\nCONFIDENTIAL")
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x8E, 0x8E, 0x93)

        doc.add_page_break()

        # Sections
        for section_key, section_title in SECTIONS:
            heading = doc.add_heading(section_title, level=1)
            for run in heading.runs:
                run.font.color.rgb = brand_rgb
                run.bold = True
                run.font.size = Pt(14)

            content = self._get_section_content(section_key)
            for line in content.split("\n"):
                line = line.strip()
                if line:
                    doc.add_paragraph(line)

        doc.save(output_path)
        logger.info("WISP DOCX exported to %s", output_path)
        return output_path

    def _export_html(self, output_path: str) -> str:
        """Fallback HTML export when PDF/DOCX libraries are not available."""
        brand_color = self.branding.get("primary_color", "#0066CC")
        app_name = self.branding.get("app_name", "Masri Digital")

        sections_html = []
        for section_key, section_title in SECTIONS:
            content = self._get_section_content(section_key)
            escaped = (content.replace("&", "&amp;").replace("<", "&lt;")
                       .replace(">", "&gt;").replace("\n", "<br>"))
            sections_html.append(
                f'<h2 style="color:{brand_color};margin-top:2em">{section_title}</h2>'
                f'<div style="margin-bottom:1em">{escaped}</div>'
            )

        html = f"""<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<title>WISP - {self.wisp.firm_name or 'Document'}</title>
<style>
  body {{ font-family: Calibri, -apple-system, sans-serif; max-width: 800px; margin: 0 auto; padding: 2em; }}
  h1 {{ color: {brand_color}; }}
  .cover {{ text-align: center; padding: 4em 0; }}
  .footer {{ color: #8E8E93; font-size: 0.85em; margin-top: 3em; border-top: 1px solid #eee; padding-top: 1em; }}
</style>
</head>
<body>
<div class="cover">
  <h1>{self.wisp.firm_name or ''}</h1>
  <h1>Written Information Security Program</h1>
  <p>Prepared by {app_name}</p>
  <p>Date: {datetime.utcnow().strftime('%B %d, %Y')}</p>
  {f'<p>Qualified Individual: {self.wisp.qi_name}{", " + self.wisp.qi_title if self.wisp.qi_title else ""}</p>' if self.wisp.qi_name else ''}
  <p style="color:#8E8E93">CONFIDENTIAL</p>
</div>
<hr>
{''.join(sections_html)}
<div class="footer">
  Prepared by {app_name} | masridigital.com | Confidential
</div>
</body>
</html>"""

        with open(output_path, "w", encoding="utf-8") as f:
            f.write(html)
        logger.info("WISP HTML fallback exported to %s", output_path)
        return output_path
